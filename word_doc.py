from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Word Documentation Automation', 0)

p = document.add_paragraph('A script for SDD documentation ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

records = (
    (1, '26', 'Sharath'),
    (2, '25', 'Bharath'),
    (3, '24', 'Bhavya')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Id'
hdr_cells[1].text = 'Age'
hdr_cells[2].text = 'Name'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')